from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import JSONResponse
from fastapi import Depends
from typing import List
import requests
import uvicorn
import pdfplumber
import os
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.text_splitter import CharacterTextSplitter
from langchain import OpenAI, VectorDBQA
from langchain.document_loaders import DirectoryLoader
# import magic
# import nltk
import openai
from pydantic import BaseModel
import io
import docx

from PIL import Image
import io

import cv2
from paddleocr import PaddleOCR, draw_ocr
import numpy as np
from langchain.document_loaders import TextLoader

from dotenv import load_dotenv
load_dotenv()


import os

# openai_key = os.environ.get("OPENAI_API_KEY")

# import pytesseract

openai.api_key = "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"


# from ExtractTable import *
# et_sess = ExtractTable(api_key="KPeK9TLm1JQLg9JlKo7cSpLaGITQQUbukm9UwzoZ")    # Replace your VALID API Key here
# print(et_sess.check_usage())                    # Validates API Key & show credits usage 

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()




# class ResponseData(BaseModel):
#     extracted_text: str
#     extracted_tables: list[list[str]]


# Configure CORS
origins = ["*"]  # Add allowed origins here
app.add_middleware(CORSMiddleware, allow_origins=origins, allow_methods=[""], allow_headers=[""])

vector_store_path = r"C:\Users\cheth\Documents\cognitus\server\local"
faiss_index_path = '/local/faiss'


def process(extracted_text):
    try:
        char_text_splitter = CharacterTextSplitter(chunk_size=2000, chunk_overlap=0)
        openAI_embeddings = OpenAIEmbeddings(openai_api_key="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")
        doc_texts = char_text_splitter.split_text(extracted_text)
        # vStore = Chroma.from_documents(doc_texts, openAI_embeddings)
        vStore = Chroma.from_texts(doc_texts, openAI_embeddings, persist_directory=vector_store_path)
        vStore.persist()
        # return ""
        # Perform additional processing with the vStore or other operations
        
    except Exception as e:
        print(e)
        # Handle exceptions here

@app.post("/ocr/image")
async def ocr(file: UploadFile = File(...)):
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang='en')
        
        image_data = await file.read()
        # image_pil = Image.open(io.BytesIO(image_data))
        # image_cv2 = cv2.cvtColor(np.array(image_pil), cv2.COLOR_RGB2BGR)
        
        result = ocr.ocr(image_data)
        
        text = ""
        for line in result:
            for item in line:
                text += item[1][0]
        process(text)    
        # loader = TextLoader(text)

        # docs = loader.load()
        
        process(text)    
        return text
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
    
from langchain.document_loaders import TextLoader



@app.post("/ocr/pdf")
async def ocr_pdf(file: UploadFile = File(...)):
    try:
        pdf = pdfplumber.open(file.file)
    
        full_text = ""
        
        for page in pdf.pages:
            page_text = page.extract_text()
            full_text += page_text
        
        pdf.close()
        # loader = TextLoader(full_text)

        # docs = loader.load()
        
        process(full_text)
        return full_text
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
        


def extract_text(docx_bytes):
    doc = docx.Document(io.BytesIO(docx_bytes))
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return "\n".join(text)

def extract_tables(docx_bytes):
    doc = docx.Document(io.BytesIO(docx_bytes))
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables

@app.post("/ocr/docx")
async def ocr_docx(file: UploadFile = File(...)):
    try:
        docx_bytes = await file.read()
        
        extracted_text = extract_text(docx_bytes)
        extracted_tables = extract_tables(docx_bytes)

        delimiter = ", "  # You can change the delimiter as needed

        result_string=""

        for i, table in enumerate(extracted_tables):


            for row in table:
                print(row)
                result_string += delimiter.join(row)
                print(result_string)

        final= ""+ extracted_text + result_string


        # with open('output.txt', 'w') as file:
        #     file.write(final)

        # loader = TextLoader(final)
        # print("loader=",loader)
        # docs = loader.load()
        # print("docs = ", docs)
        # process(docs)

        process(final)
        return final

    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.post("/api")
async def call_api(text: str = Form(...)):
    try:
        # endpoint_url = 'https://ap-south-1.console.aws.amazon.com/sagemaker/home?region=ap-south-1#/endpoints/huggingface-pytorch-tgi-inference-2023-08-16-12-04-56-323/invocations'
        
        # response = requests.post(endpoint_url, json={
        #     'inputs': text
        # }).json()
        print(text)
        openAI_embeddings = OpenAIEmbeddings(openai_api_key="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")
        # return text
        vStore = Chroma(embedding_function=openAI_embeddings,persist_directory= vector_store_path)
        vStore.get()
        model = VectorDBQA.from_chain_type(llm=OpenAI(model_name="text-davinci-003"), chain_type="stuff", vectorstore=vStore) 
        response = model.run(text)
        

        print(response)
        return response
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)
        # print('There seems to be an issue. We will get back to you soon: ', e)



if __name__ == '__main__':
    uvicorn.run("main:app", host="0.0.0.0", port=5000, reload=True)